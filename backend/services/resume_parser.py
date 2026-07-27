import io
import magic
from typing import Tuple, Optional, Tuple

import pdfplumber
from docx import Document
import PyPDF2

from backend.utils.file_utils import (
    FileParsingError,
    TextExtractionError,
    FileUploadError,
    log_error,
    log_warning,
    log_info,
    with_fallback
)


from backend.core.config import (
    MAX_FILE_SIZE_BYTES,
    MAX_FILE_SIZE_MB,
    SUPPORTED_MIME_TYPES,
)



class FileParsingError(Exception):
    """Custom exception for file parsing errors."""
    pass

class FileValidationError(Exception):
    """Custom exception for file validation errors."""
    pass




def validate_file(file_data:bytes, filename:str) -> Tuple[bool, str, Optional[str]]:
    """
    Validates the uploaded file for size and type.
    Raises FileValidationError if validation fails.
    """
    file_size_bytes = len(file_data)
    if file_size_bytes > MAX_FILE_SIZE_BYTES:
        size_mb = file_size_bytes / (1024 * 1024)
        return False, ( f'File size ({size_mb:.2f} MB) exceeds the maximum of {MAX_FILE_SIZE_MB} MB. '
            'Please upload a smaller file or compress your resume.'), None
    if file_size_bytes == 0:
        return False, 'uploaded file is empty...please check the file you have uploaded and try again', None

    try:
        mime_type = magic.from_buffer(file_data, mime=True)#check the file content type using magic library

    except Exception as e:
        return False, f"error determining the file type : {e}", None

    if mime_type not in SUPPORTED_MIME_TYPES:
        supported=', '.join(SUPPORTED_MIME_TYPES.keys()).upper()
        return False, (
            f'Unsupported file type: {mime_type}. '
            f'Please upload one of: {supported}.'
        ), None
    return True, '', SUPPORTED_MIME_TYPES[mime_type]


def _extract_pdf_hyperlinks(file_data: bytes) -> str:
    urls = []
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_data))
        for page in reader.pages:
            if '/Annots' not in page: # if not present hyper links, skip to next page
                continue
            for annot_ref in page['/Annots']: # if page has multiple annotations, iterate through them
                try:
                    annot = annot_ref.get_object() # get the actual annotation object
                    if annot.get('/Subtype') != '/Link': # if the annotation is not a link, skip to next annotation
                        continue
                    action = annot.get('/A', {})# get the action dictionary from the annotation
                    uri = action.get('/URI', '') # get the URI(Uniform Resource Identifier) from the action dictionary
                    if uri and isinstance(uri, (str, bytes)): # check if the URI is not empty and is either a string or bytes becaude pyPDF2 may read bytes and string 
                        # PyPDF2 may return bytes for URI values
                        if isinstance(uri, bytes):
                            uri = uri.decode('utf-8', errors='ignore')  #decode bytes to string, somtime uri looks like b'http://example.com' so we need to decode it to string
                        uri = uri.strip() #strip any leading or trailing whitespace from the URI
                        if uri.startswith('http'): # check if the URI starts with 'http' to ensure it's a valid URL
                            urls.append(uri) # add the valid URL to the list of URLs
                except Exception:
                    pass
    except Exception:
        pass
    return '\n'.join(urls)



def _extract_pdf_with_pdfplumber(file_data: bytes) -> str:
    text = ''  # Initialize an empty string to store the extracted text.

    # Open the PDF file from memory using pdfplumber.
    with pdfplumber.open(io.BytesIO(file_data)) as pdf:

        # Iterate through each page in the PDF.
        for page in pdf.pages:

            # Extract text from the current page.
            page_text = page.extract_text()

            # If text was extracted, append it to the final text.
            if page_text:
                text += page_text + '\n'

    # If no text was extracted, raise a custom exception.
    if not text.strip():
        raise TextExtractionError(
            'pdfplumber extracted no text',
            user_message='No text could be extracted from the PDF.'
        )

    # Extract clickable hyperlinks from the PDF.
    hyperlinks = _extract_pdf_hyperlinks(file_data)

    # If hyperlinks exist, append them to the extracted text.
    if hyperlinks:
        text = text.strip() + '\n' + hyperlinks

    # Return the extracted text with hyperlinks.
    return text.strip()


def _extract_pdf_with_pypdf2(file_data: bytes) -> str:
    # Initialize an empty string to store the extracted text.
    text = ''

    # Read the PDF from memory using PyPDF2.
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))

    # Iterate through each page in the PDF.
    for page in pdf_reader.pages:

        # Extract text from the current page.
        page_text = page.extract_text()

        # If text was extracted, append it to the final text.
        if page_text:
            text += page_text + '\n'

    # If no text was extracted, raise a custom exception.
    if not text.strip():
        raise TextExtractionError(
            'PyPDF2 extracted no text',
            user_message='No text could be extracted from the PDF.'
        )

    # Extract clickable hyperlinks from the PDF.
    hyperlinks = _extract_pdf_hyperlinks(file_data)

    # If hyperlinks exist, append them to the extracted text.
    if hyperlinks:
        text = text.strip() + '\n' + hyperlinks

    # Return the extracted text with hyperlinks.
    return text.strip()




def extract_text_from_pdf(file_data: bytes) -> str:
    # Try to extract text from the PDF.
    try:

        # First, use pdfplumber.
        # If it fails, automatically use PyPDF2 as a backup.
        result, used_fallback = with_fallback(
            _extract_pdf_with_pdfplumber,
            _extract_pdf_with_pypdf2,
            file_data,
            log_fallback=True
        )

        # Check if the backup (PyPDF2) was used.
        if used_fallback:

            # Log that the PDF was successfully extracted using PyPDF2.
            log_info(
                'PDF EXTRACTION succeeded using the PyPDF2 fallback',
                context='resume_parser'
            )

        # Return the extracted text.
        return result

    # If both methods fail, handle the error.
    except Exception as e:

        # Log the error for debugging.
        log_error(e, context='extract_text_from_pdf')

        # Raise a custom error with a user-friendly message.
        raise FileParsingError(
            'Failed to extract text from PDF using both pdfplumber and PyPDF2. '
            'The PDF may be corrupted, password-protected, or contain only scanned images. '
            'Please ensure it contains selectable text.'
        ) from e  # Keep the original error information.



def extract_text_from_docx(file_data: bytes) -> str:
    # Try to extract text from the DOCX file.
    try:

        # Open the DOCX file from memory.
        doc = Document(io.BytesIO(file_data))

        # Create an empty list to store the extracted text.
        text_parts = []

        # Loop through all paragraphs in the document.
        for paragraph in doc.paragraphs:

            # If the paragraph is not empty, add it to the list.
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Loop through all tables in the document.
        for table in doc.tables:

            # Loop through each row in the table.
            for row in table.rows:

                # Loop through each cell in the row.
                for cell in row.cells:

                    # If the cell contains text, add it to the list.
                    if cell.text.strip():
                        text_parts.append(cell.text)

        # Join all extracted text into one string.
        text = '\n'.join(text_parts)

        # If no text was extracted, raise a custom error.
        if not text.strip():
            raise FileParsingError(
                'No text could be extracted from the document. '
                'The document may be empty or corrupted.'
            )

        # Try to extract hyperlinks from the document.
        try:

            # Loop through all relationships in the DOCX file.
            for rel in doc.part.rels.values():

                # Check if the relationship is a hyperlink.
                if 'hyperlink' in rel.reltype.lower():

                    # Get the hyperlink URL.
                    url = rel._target

                    # If the URL is valid, add it to the extracted text.
                    if isinstance(url, str) and url.startswith('http'):
                        text += '\n' + url

        # Ignore errors while extracting hyperlinks.
        except Exception:
            pass

        # Log how many characters were extracted.
        log_info(f'Extracted {len(text)} chars from DOCX', context='resume_parser')

        # Return the extracted text.
        return text.strip()

    # If a FileParsingError was already raised, raise it again unchanged.
    except FileParsingError:
        raise

    # Handle any other unexpected errors.
    except Exception as e:

        # Log the error for debugging.
        log_error(e, context='extract_text_from_docx')

        # Raise a user-friendly custom error.
        raise FileParsingError(
            'Failed to extract text from DOCX. '
            'The document may be corrupted or in an unsupported format. '
            'Please try re-saving or converting to PDF.'
        ) from e  # Keep the original error information.


def extract_text_from_doc(file_data: bytes) -> str:
    # .doc (old Microsoft Word format) is not supported.
    # Raise a custom error and tell the user to convert the file.
    raise FileParsingError(
        'Legacy .doc format is not supported. '
        'Please convert your document to .docx or .pdf and try again. '
        'You can convert using Microsoft Word, Google Docs, or online tools.'
    )


def extract_text(file_data: bytes, file_type: str) -> str:
    # Check if the file is a PDF.
    if file_type == 'pdf':

        # Extract text from the PDF.
        return extract_text_from_pdf(file_data)

    # Check if the file is a DOCX.
    elif file_type == 'docx':

        # Extract text from the DOCX file.
        return extract_text_from_docx(file_data)

    # Check if the file is an old DOC file.
    elif file_type == 'doc':

        # Raise an error because .doc files are not supported.
        return extract_text_from_doc(file_data)

    # If the file type is not supported.
    else:

        # Raise a custom validation error.
        raise FileValidationError(
            f'Invalid file type: {file_type}. '
            'Supported types are: pdf, docx, and doc.'
        )




def parse_resume_file(file_data: bytes, filename: str) -> Tuple[str, dict]:
    # Log that resume parsing has started.
    log_info(f'Parsing file: {filename}', context='parse_resume_file')

    # ---------------------------
    # Phase 1: Validate the file
    # ---------------------------
    try:
        # Check if the uploaded file is valid and get its type.
        is_valid, error_msg, file_type = validate_file(file_data, filename)

        # If the file is not valid, raise a validation error.
        if not is_valid:
            log_warning(
                f'Validation failed for file {filename}',
                context='parse_resume_file'
            )
            raise FileValidationError(error_msg)

    # If a FileValidationError was already raised, pass it unchanged.
    except FileValidationError:
        raise

    # Handle any unexpected errors during validation.
    except Exception as e:
        # Log the error for debugging.
        log_error(e, context='parse_resume_file_validation')

        # Raise a user-friendly validation error.
        raise FileValidationError(
            'Could not validate the uploaded file. '
            'Please ensure it is a valid PDF or DOCX.'
        ) from e

    # ---------------------------------
    # Phase 2: Extract text from the file
    # ---------------------------------
    try:
        # Extract text based on the file type.
        text = extract_text(file_data, file_type)

        # Log the number of characters extracted.
        log_info(
            f'Extracted {len(text)} chars from {filename}',
            context='parse_resume_file'
        )

    # If a FileParsingError was already raised, pass it unchanged.
    except FileParsingError:
        raise

    # Handle any unexpected errors during text extraction.
    except Exception as e:
        # Log the error for debugging.
        log_error(e, context='parse_resume_file_extraction')

        # Raise a user-friendly parsing error.
        raise FileParsingError(
            'An unexpected error occurred while processing the file. '
            'Please try again or contact support if the problem persists.'
        ) from e

    # ---------------------------
    # Phase 3: Create file metadata
    # ---------------------------

    # Store information about the processed file.
    metadata = {
        'filename': filename,                  # Original file name
        'file_type': file_type,                # pdf / docx
        'file_size_bytes': len(file_data),     # File size in bytes
        'text_length': len(text),              # Number of extracted characters
        'success': True,                       # Parsing completed successfully
    }

    # Return the extracted text and metadata.
    return text, metadata